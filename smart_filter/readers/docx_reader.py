from __future__ import annotations

from pathlib import Path


def _clean_cells(values: list[object]) -> list[str]:
    return [str(value or "").strip() for value in values]


def _looks_like_table_header(row: list[str], following_rows: list[list[str]]) -> bool:
    nonempty = [value for value in row if value]
    if len(nonempty) < 2 or not following_rows:
        return False
    if len(set(value.casefold() for value in nonempty)) != len(nonempty):
        return False
    return all(len(value) <= 60 and any(character.isalpha() for character in value) for value in nonempty)


def _table_lines(table) -> list[str]:
    rows = [_clean_cells([cell.text for cell in row.cells]) for row in table.rows]
    if not rows:
        return []

    output: list[str] = []
    headers = rows[0]
    if _looks_like_table_header(headers, rows[1:]):
        for row in rows[1:]:
            emitted = False
            for index, value in enumerate(row):
                if not value:
                    continue
                header = headers[index].strip() if index < len(headers) else ""
                output.append(f"{header}: {value}" if header else value)
                emitted = True
            if emitted:
                output.append("")
        return output

    for row in rows:
        values = [value for value in row if value]
        if values:
            output.append(" | ".join(values))
    return output


def read_docx_text(file_path: str | Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on local environment.
        raise RuntimeError("Falta dependencia python-docx para leer archivos .docx") from exc

    extracted_text: list[str] = []
    document = Document(str(file_path))

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if paragraph_text:
            extracted_text.append(paragraph_text)

    for table in document.tables:
        extracted_text.extend(_table_lines(table))

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                extracted_text.append(paragraph_text)
        for paragraph in section.footer.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                extracted_text.append(paragraph_text)

    return "\n".join(extracted_text)
