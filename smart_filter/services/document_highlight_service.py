from __future__ import annotations

import csv
import json
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from date_time_core import create_timestamp_pair
from render_core import render_report_data

from smart_filter.app_info import APP_NAME, APP_VERSION
from smart_filter.domain.search_models import SearchResult
from smart_filter.domain.text_normalizer import normalize_text
from smart_filter.readers.csv_reader import detect_csv_dialect, read_text_with_fallback as read_csv_source
from smart_filter.readers.reader_registry import read_file_content
from smart_filter.readers.text_reader import read_text_with_fallback
from smart_filter.services.settings_service import get_settings

MAX_TEXT_CHARS = 8_000_000
MAX_TABLE_ROWS = 25_000
MAX_TABLE_CELLS = 250_000


def _format_file_size(size_bytes: int) -> str:
    value = float(max(0, int(size_bytes)))
    units = ("B", "KB", "MB", "GB", "TB")
    unit = units[0]
    for candidate in units:
        unit = candidate
        if value < 1024 or candidate == units[-1]:
            break
        value /= 1024
    if unit == "B":
        return f"{int(value)} {unit}"
    return f"{value:.1f} {unit}"


@dataclass(frozen=True)
class DocumentHighlightRender:
    output_path: Path
    sections_count: int
    terms_count: int
    truncated: bool = False
    diagnostics: tuple[str, ...] = ()


class _RenderBudget:
    def __init__(self) -> None:
        self.text_chars = 0
        self.table_rows = 0
        self.table_cells = 0
        self.truncated = False

    def allow_text(self, text: str) -> str:
        value = str(text or "")
        remaining = MAX_TEXT_CHARS - self.text_chars
        if remaining <= 0:
            self.truncated = True
            return ""
        if len(value) > remaining:
            self.truncated = True
            value = value[:remaining]
        self.text_chars += len(value)
        return value

    def allow_row(self, cells: list[str]) -> bool:
        if self.table_rows >= MAX_TABLE_ROWS or self.table_cells + len(cells) > MAX_TABLE_CELLS:
            self.truncated = True
            return False
        self.table_rows += 1
        self.table_cells += len(cells)
        return True


def _clean_terms(values: Iterable[object]) -> list[str]:
    terms: list[str] = []
    seen: set[str] = set()
    for value in values:
        clean = str(value or "").strip()
        key = normalize_text(clean)
        if not clean or not key or key in seen:
            continue
        seen.add(key)
        terms.append(clean)
    return terms


def _line_blocks(
    text: str,
    *,
    budget: _RenderBudget,
    line_start: int = 1,
    location_prefix: str = "",
) -> tuple[list[dict[str, Any]], int]:
    safe_text = budget.allow_text(text)
    lines = safe_text.splitlines()
    if not lines and safe_text:
        lines = [safe_text]
    blocks: list[dict[str, Any]] = []
    line_number = line_start
    for raw_line in lines:
        location = f"{location_prefix} · Línea {line_number}" if location_prefix else f"Línea {line_number}"
        blocks.append(
            {
                "type": "line",
                "line_number": line_number,
                "location_label": location,
                "text": raw_line,
            }
        )
        line_number += 1
    return blocks, line_number


def _text_section(
    label: str,
    text: str,
    *,
    budget: _RenderBudget,
    section_id: str = "content",
    kind: str = "texto",
) -> list[dict[str, Any]]:
    blocks, _ = _line_blocks(text, budget=budget)
    return [{"id": section_id, "label": label, "kind": kind, "blocks": blocks}]


def _pdf_sections(source: Path, budget: _RenderBudget) -> list[dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on runtime.
        raise RuntimeError("Falta dependencia pypdf para generar la vista PDF") from exc

    reader = PdfReader(str(source))
    if getattr(reader, "is_encrypted", False) and reader.decrypt("") == 0:
        raise ValueError("PDF protegido o cifrado")

    sections: list[dict[str, Any]] = []
    global_line = 1
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        blocks, global_line = _line_blocks(
            text,
            budget=budget,
            line_start=global_line,
            location_prefix=f"Página {page_number}",
        )
        sections.append(
            {
                "id": f"page-{page_number}",
                "label": f"Página {page_number}",
                "kind": "página PDF",
                "blocks": blocks,
            }
        )
        if budget.truncated:
            break
    return sections


def _docx_sections(source: Path, budget: _RenderBudget) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on runtime.
        raise RuntimeError("Falta dependencia python-docx para generar la vista Word") from exc

    document = Document(str(source))
    blocks: list[dict[str, Any]] = []
    paragraph_number = 0
    for paragraph in document.paragraphs:
        text = budget.allow_text(paragraph.text)
        if not text and not paragraph.text:
            continue
        paragraph_number += 1
        blocks.append(
            {
                "type": "paragraph",
                "location_label": f"Párrafo {paragraph_number}",
                "text": text,
            }
        )
        if budget.truncated:
            break

    for table_number, table in enumerate(document.tables, start=1):
        rows: list[dict[str, Any]] = []
        max_columns = 0
        for row_number, row in enumerate(table.rows, start=1):
            cells = [budget.allow_text(cell.text) for cell in row.cells]
            if not budget.allow_row(cells):
                break
            max_columns = max(max_columns, len(cells))
            rows.append(
                {
                    "number": row_number,
                    "location_label": f"Tabla {table_number} · Fila {row_number}",
                    "cells": cells,
                }
            )
        blocks.append(
            {
                "type": "table",
                "location_label": f"Tabla {table_number}",
                "headers": [],
                "max_columns": max_columns,
                "rows": rows,
            }
        )
        if budget.truncated:
            break

    return [{"id": "document", "label": "Documento", "kind": "Word", "blocks": blocks}]


def _xlsx_sections(source: Path, budget: _RenderBudget) -> list[dict[str, Any]]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - depends on runtime.
        raise RuntimeError("Falta dependencia openpyxl para generar la vista Excel") from exc

    workbook = load_workbook(source, read_only=True, data_only=True)
    sections: list[dict[str, Any]] = []
    try:
        for sheet_index, worksheet in enumerate(workbook.worksheets, start=1):
            rows: list[dict[str, Any]] = []
            max_columns = 0
            for row_number, values in enumerate(worksheet.iter_rows(values_only=True), start=1):
                cells = [budget.allow_text("" if value is None else str(value)) for value in values]
                while cells and not cells[-1]:
                    cells.pop()
                if not cells:
                    continue
                if not budget.allow_row(cells):
                    break
                max_columns = max(max_columns, len(cells))
                rows.append(
                    {
                        "number": row_number,
                        "location_label": f"Hoja {worksheet.title} · Fila {row_number}",
                        "cells": cells,
                    }
                )
            sections.append(
                {
                    "id": f"sheet-{sheet_index}",
                    "label": f"Hoja · {worksheet.title}",
                    "kind": "hoja Excel",
                    "blocks": [
                        {
                            "type": "table",
                            "location_label": f"Hoja {worksheet.title}",
                            "headers": [],
                            "max_columns": max_columns,
                            "rows": rows,
                        }
                    ],
                }
            )
            if budget.truncated:
                break
    finally:
        workbook.close()
    return sections


def _csv_sections(source: Path, budget: _RenderBudget) -> list[dict[str, Any]]:
    raw = read_csv_source(source)
    dialect = detect_csv_dialect(raw)
    rows: list[dict[str, Any]] = []
    max_columns = 0
    for row_number, values in enumerate(csv.reader(raw.splitlines(), dialect=dialect), start=1):
        cells = [budget.allow_text(value) for value in values]
        if not budget.allow_row(cells):
            break
        max_columns = max(max_columns, len(cells))
        rows.append({"number": row_number, "location_label": f"Fila {row_number}", "cells": cells})
    return [
        {
            "id": "csv-table",
            "label": "Contenido CSV",
            "kind": "tabla CSV",
            "blocks": [
                {
                    "type": "table",
                    "location_label": "CSV",
                    "headers": [],
                    "max_columns": max_columns,
                    "rows": rows,
                }
            ],
        }
    ]


def _pretty_json(source: Path) -> str:
    raw = read_text_with_fallback(source)
    try:
        return json.dumps(json.loads(raw), ensure_ascii=False, indent=2)
    except Exception:
        return raw


def _pretty_xml(source: Path) -> str:
    raw = read_text_with_fallback(source)
    try:
        root = ET.fromstring(raw)
        ET.indent(root, space="  ")
        return ET.tostring(root, encoding="unicode")
    except Exception:
        return raw


def _generic_text(source: Path) -> tuple[str, str]:
    extension = source.suffix.lower()
    if extension == ".json":
        return _pretty_json(source), "JSON"
    if extension == ".xml":
        return _pretty_xml(source), "XML"
    if extension in {".txt", ".log", ".md"}:
        return read_text_with_fallback(source), extension.lstrip(".").upper() or "Texto"

    reader_result = read_file_content(source)
    if reader_result.status != "ok":
        raise RuntimeError(reader_result.error or f"No se pudo leer {source.name}")
    return reader_result.text, reader_result.reader_name


def _build_sections(source: Path) -> tuple[list[dict[str, Any]], _RenderBudget, str]:
    budget = _RenderBudget()
    extension = source.suffix.lower()
    if extension == ".pdf":
        return _pdf_sections(source, budget), budget, "pdf_reader"
    if extension == ".docx":
        return _docx_sections(source, budget), budget, "docx_reader"
    if extension == ".xlsx":
        return _xlsx_sections(source, budget), budget, "xlsx_reader"
    if extension == ".csv":
        return _csv_sections(source, budget), budget, "csv_reader"

    text, reader_name = _generic_text(source)
    return _text_section("Contenido", text, budget=budget, kind=reader_name), budget, reader_name


def _highlight_locations(result: SearchResult) -> list[dict[str, Any]]:
    locations = getattr(result, "match_locations", None)
    if isinstance(locations, list) and locations:
        return [dict(item) for item in locations if isinstance(item, dict)]
    return [
        {
            "location_label": result.location_label or "Archivo",
            "line_number": result.line_number,
            "row_number": result.row_number,
            "sheet_name": result.sheet_name,
            "matched_terms": list(result.matched_terms),
            "preview_text": result.preview_text,
        }
    ]


def build_document_highlight_contract(result: SearchResult) -> tuple[dict[str, Any], tuple[str, ...]]:
    source = Path(result.full_path).expanduser().resolve()
    if not source.exists() or not source.is_file():
        raise FileNotFoundError(f"El archivo original ya no existe: {source}")

    sections, budget, reader_name = _build_sections(source)
    terms = _clean_terms(result.matched_terms)
    if not terms:
        terms = _clean_terms([result.matches])
    if not terms:
        raise ValueError("No hay términos disponibles para destacar")

    generated_at_utc, generated_at_local = create_timestamp_pair()
    diagnostics: list[dict[str, Any]] = []
    diagnostic_messages: list[str] = []
    if budget.truncated:
        message = (
            "La vista HTML fue limitada por seguridad; el archivo original permanece completo y sin modificaciones."
        )
        diagnostics.append({"level": "warning", "code": "document_view_truncated", "message": message})
        diagnostic_messages.append(message)

    locations = _highlight_locations(result)
    occurrence_count = int(getattr(result, "occurrence_count", 0) or len(locations) or 1)
    contract = {
        "meta": {
            "schema_version": "1.0.0",
            "tool_name": APP_NAME,
            "tool_version": APP_VERSION,
            "version": APP_VERSION,
            "report_type": "document_highlight",
            "generated_at": generated_at_local,
            "generated_at_utc": generated_at_utc,
            "source_format": source.suffix.lower(),
        },
        "summary": {
            "status": "warning" if budget.truncated else "ok",
            "match_occurrences_count": occurrence_count,
            "matched_terms_count": len(terms),
            "sections_count": len(sections),
            "truncated": budget.truncated,
        },
        "report_brief": {
            "title": source.name,
            "subtitle": f"Vista HTML destacada · {result.category_name}",
            "status": "warning" if budget.truncated else "ok",
            "description": "Representación HTML temporal generada por Smart Filter con RenderCore.",
        },
        "data": {
            "document": {
                "title": source.name,
                "source_path": str(source),
                "source_uri": source.as_uri(),
                "source_directory_uri": source.parent.as_uri(),
                "source_size_bytes": source.stat().st_size,
                "source_size_label": _format_file_size(source.stat().st_size),
                "format": source.suffix.lower().lstrip(".") or "archivo",
                "reader": reader_name,
                "truncated": budget.truncated,
                "sections": sections,
            },
            "highlight": {
                "category_name": result.category_name,
                "terms": [{"text": term} for term in terms],
                "occurrence_count": occurrence_count,
                "locations": locations,
            },
        },
        "diagnostics": diagnostics,
        "errors": [],
    }
    return contract, tuple(diagnostic_messages)


def render_document_highlight(result: SearchResult, output_path: str | Path) -> DocumentHighlightRender:
    contract, diagnostic_messages = build_document_highlight_contract(result)
    settings = get_settings()
    appearance_mode = str(settings.get("appearance_mode") or settings.get("theme") or "dark").strip().lower()
    theme = "light" if appearance_mode in {"light", "claro"} else "dark"
    output = Path(output_path).expanduser().resolve()
    render_result = render_report_data(
        contract,
        output,
        output_format="html",
        profile="document_highlight_pro",
        theme=theme,
        contract_profile="tool_report",
        input_name="smartfilter_document_highlight.json",
    )
    if not render_result.ok:
        raise RuntimeError(render_result.message)
    document = contract["data"]["document"]
    return DocumentHighlightRender(
        output_path=output,
        sections_count=len(document.get("sections", [])),
        terms_count=len(contract["data"]["highlight"].get("terms", [])),
        truncated=bool(document.get("truncated", False)),
        diagnostics=diagnostic_messages,
    )
