from __future__ import annotations

import json
import tempfile
from pathlib import Path

from smart_filter.bootstrap import ensure_sharedcode_on_path

ensure_sharedcode_on_path()

from smart_filter.engine.category_scope import extract_category_content_scope
from smart_filter.readers.csv_reader import read_csv_text
from smart_filter.readers.docx_reader import read_docx_text
from smart_filter.readers.text_reader import read_json_text, read_xml_text
from smart_filter.readers.xlsx_reader import read_xlsx_text


def _assert_scope(content: str, target: str, expected: str) -> None:
    scope = extract_category_content_scope(content, [target])
    assert scope.text == expected, (target, scope.text, expected)


def main() -> int:
    # Una frase normal no puede transformarse en sección por comenzar con el
    # nombre configurado.
    _assert_scope(
        "Experiencia administrativa requerida para el puesto.",
        "Experiencia",
        "",
    )
    _assert_scope(
        "La experiencia del usuario fue satisfactoria y quedó documentada.",
        "Experiencia",
        "",
    )

    # Encabezados reales, incluso generales, abren y cierran bloques.
    _assert_scope(
        "\n".join(
            [
                "1. INCIDENCIAS",
                "Error de permisos en administración.",
                "El servicio quedó detenido.",
                "2. ACCIONES",
                "Reinicio controlado.",
            ]
        ),
        "Incidencias",
        "Error de permisos en administración.\nEl servicio quedó detenido.",
    )
    _assert_scope(
        "## Resultados\nAdministración completada\n\n## Conclusiones\nSin cambios",
        "Resultados",
        "Administración completada",
    )

    # Etiqueta/valor toma solamente el valor asociado, no todo lo posterior.
    _assert_scope(
        "Experiencia: Administración de servidores\nResponsable: Ana\nEstado: Cerrado",
        "Experiencia",
        "Administración de servidores",
    )

    # Compatibilidad con CV aplanado: solo se habilita cuando existen varias
    # etiquetas conocidas que prueban una estructura real.
    _assert_scope(
        "Curriculum Vitae Nombre Ana Perfil Soporte Experiencia Administración Facturación Habilidades Excel",
        "Experiencia",
        "Administración Facturación",
    )

    with tempfile.TemporaryDirectory(prefix="smartfilter_section_detection_") as temp_dir:
        root = Path(temp_dir)

        csv_path = root / "personas.csv"
        csv_path.write_text(
            "Nombre;Experiencia;Educacion\nAna;Administración de servidores;Curso técnico\n",
            encoding="utf-8",
        )
        csv_text = read_csv_text(csv_path)
        assert "Experiencia: Administración de servidores" in csv_text
        _assert_scope(csv_text, "Experiencia", "Administración de servidores")

        json_path = root / "reporte.json"
        json_path.write_text(
            json.dumps(
                {
                    "incidencias": {
                        "detalle": "Fallo en administración de permisos",
                        "estado": "abierta",
                    },
                    "acciones": {"detalle": "Revisar grupos"},
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        json_text = read_json_text(json_path)
        json_scope = extract_category_content_scope(json_text, ["Incidencias"])
        assert "Fallo en administración de permisos" in json_scope.text
        assert "Revisar grupos" not in json_scope.text

        xml_path = root / "reporte.xml"
        xml_path.write_text(
            "<reporte><incidencias><detalle>Error administrativo</detalle></incidencias>"
            "<acciones><detalle>Reiniciar</detalle></acciones></reporte>",
            encoding="utf-8",
        )
        xml_text = read_xml_text(xml_path)
        xml_scope = extract_category_content_scope(xml_text, ["Incidencias"])
        assert "Error administrativo" in xml_scope.text
        assert "Reiniciar" not in xml_scope.text

        try:
            from openpyxl import Workbook
        except ImportError as exc:  # pragma: no cover
            raise AssertionError("openpyxl es obligatorio para esta validación") from exc

        xlsx_path = root / "personas.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Datos"
        sheet.append(["Nombre", "Experiencia", "Educacion"])
        sheet.append(["Ana", "Administración de servidores", "Curso técnico"])
        workbook.save(xlsx_path)
        workbook.close()
        xlsx_text = read_xlsx_text(xlsx_path)
        assert "Experiencia: Administración de servidores" in xlsx_text
        _assert_scope(xlsx_text, "Experiencia", "Administración de servidores")

        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise AssertionError("python-docx es obligatorio para esta validación") from exc

        docx_path = root / "informe.docx"
        document = Document()
        document.add_paragraph("INCIDENCIAS")
        document.add_paragraph("Fallo en administración de usuarios")
        document.add_paragraph("ACCIONES")
        document.add_paragraph("Revisar permisos")
        document.save(docx_path)
        docx_text = read_docx_text(docx_path)
        _assert_scope(docx_text, "Incidencias", "Fallo en administración de usuarios")

    print("CATEGORY_SECTION_DETECTION_OK")
    print(
        {
            "casual_sentence": "ignored",
            "general_headings": "ok",
            "structured_readers": ["csv", "xlsx", "docx", "json", "xml"],
            "legacy_flat_cv": "compatible",
        }
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
